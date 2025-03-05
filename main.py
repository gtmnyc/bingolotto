from kivy.config import Config
Config.set('graphics', 'width', '360')
Config.set('graphics', 'height', '640')

import os
import json
import dropbox
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from kivy.app import App
from kivy.lang import Builder
from kivy.properties import ListProperty, StringProperty, BooleanProperty
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.scrollview import ScrollView
from kivy.uix.popup import Popup
from kivy.uix.label import Label
from kivy.uix.gridlayout import GridLayout
from kivy.uix.button import Button
from kivy.utils import platform

###########################
# Dropbox config
###########################
APP_KEY = '99hl0e4g22uysd1'
APP_SECRET = 'w16oh618rf7u56i'
REFRESH_TOKEN = 'IJrcOdm637AAAAAAAAAAAfpY8pX22gQCwy03vc_Cg6L5m4x4yBYovmjk7aQRbQSP'
DROPBOX_FILE_PATH = '/Apps/bingo lotto/Bingo-Lotto-board.docx'
LOCAL_FILE_PATH = 'Bingo-Lotto-board.docx'

def get_dropbox_client():
    return dropbox.Dropbox(
        oauth2_refresh_token=REFRESH_TOKEN,
        app_key=APP_KEY,
        app_secret=APP_SECRET
    )

def download_file():
    dbx = get_dropbox_client()
    if os.path.exists(LOCAL_FILE_PATH):
        try:
            os.remove(LOCAL_FILE_PATH)
        except PermissionError:
            pass
    with open(LOCAL_FILE_PATH, 'wb') as f:
        _, res = dbx.files_download(path=DROPBOX_FILE_PATH)
        f.write(res.content)

def upload_file():
    dbx = get_dropbox_client()
    with open(LOCAL_FILE_PATH, 'rb') as f:
        dbx.files_upload(
            f.read(),
            DROPBOX_FILE_PATH,
            mode=dropbox.files.WriteMode('overwrite')
        )

###########################
# Winner logic
###########################
def highlight_cell(cell, color):
    tc_pr = cell._element.get_or_add_tcPr()
    cell_shading = OxmlElement('w:shd')
    cell_shading.set(qn('w:val'), 'clear')
    cell_shading.set(qn('w:color'), 'auto')
    cell_shading.set(qn('w:fill'), color)
    tc_pr.append(cell_shading)

def clear_cell_highlight(cell):
    tc_pr = cell._element.get_or_add_tcPr()
    for shading in tc_pr.findall(qn('w:shd')):
        tc_pr.remove(shading)

def clear_date_table(table):
    while len(table.rows) > 1:
        table._element.remove(table.rows[-1]._tr)

def delete_date_table(doc):
    for table in doc.tables:
        if table.cell(0, 0).text == "Date":
            table._element.getparent().remove(table._element)
            break

###########################
# Force black borders on each cell
###########################
def set_cell_borders(cell):
    """
    Force black single-line borders on top/left/bottom/right
    of this specific cell.
    """
    tc_pr = cell._element.get_or_add_tcPr()

    # Remove existing <w:tcBorders> if any
    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is not None:
        tc_pr.remove(borders)

    # Create new <w:tcBorders>
    borders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'single')
        edge_el.set(qn('w:sz'), '8')        # thickness
        edge_el.set(qn('w:space'), '0')
        edge_el.set(qn('w:color'), '000000')  # black
        borders.append(edge_el)
    tc_pr.append(borders)

def add_numbers_to_date_table(doc, numbers):
    """
    Re-creates or updates the 'Date' table.
    Then for each cell (header + newly added row), force black borders.
    """
    date_table = None
    for table in doc.tables:
        if table.cell(0, 0).text == 'Date':
            date_table = table
            break

    if date_table is None:
        date_table = doc.add_table(rows=1, cols=8)
        hdr = date_table.rows[0].cells
        hdr[0].text = 'Date'
        for i in range(1, 8):
            hdr[i].text = f'Ball {i}'
        for c in hdr:
            for paragraph in c.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for c in hdr:
            set_cell_borders(c)

    row_cells = date_table.add_row().cells
    row_cells[0].text = datetime.now().strftime('%d-%m-%y')
    for i, val in enumerate(numbers[:7]):
        row_cells[i+1].text = str(val)
    for c in row_cells:
        for paragraph in c.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_borders(c)

    for row in date_table.rows:
        for cell in row.cells:
            set_cell_borders(cell)

def clear_highlights(doc, cumulative_matches, highlight_sequence):
    for table in doc.tables:
        if table.cell(0, 0).text in ["Date", "METADATA"]:
            continue
        for row in table.rows:
            for cell in row.cells:
                clear_cell_highlight(cell)

    cumulative_matches.clear()
    highlight_sequence.clear()

    for t in doc.tables:
        if t.cell(0, 0).text == "METADATA":
            t._element.getparent().remove(t._element)

    delete_date_table(doc)

def calculate_jackpot(doc, num_entrants=110, entry_fee=5):
    date_table = None
    for t in doc.tables:
        if t.cell(0, 0).text == 'Date':
            date_table = t
            break
    if date_table:
        total_weeks = len(date_table.rows)
        return num_entrants * entry_fee * total_weeks
    return 0

def generate_winner_message(winners, close_calls, jackpot):
    if not winners:
        return "No winner yet."
    msg = f"""
    **Congratulations to our Winners!**

    A big congratulations to our winners this week: **{', '.join(winners)}**.
    You've done it—matching all six numbers to take home the prize of {jackpot}!

    A special mention to **{', '.join(close_calls)}**.
    You were so close, matching all six numbers but just missed out
    as our winners matched all six first.

    Good luck to all next week!
    """
    return msg.strip()

def highlight_number(doc, numbers, cumulative_matches, highlight_sequence):
    console_lines = []
    matched_names_5 = set()
    matched_names_6 = set()
    match_order = {}

    for num in numbers:
        for table in doc.tables:
            if table.cell(0, 0).text in ["METADATA", "Date"]:
                continue
            for row in table.rows:
                if len(row.cells) < 2:
                    continue
                name_cell = row.cells[1]
                name = name_cell.text.strip()

                if name not in cumulative_matches:
                    cumulative_matches[name] = set()
                if name not in match_order:
                    match_order[name] = []

                for c in row.cells[2:]:
                    if c.text.strip() == num:
                        highlight_cell(c, 'FFB6C1')
                        cumulative_matches[name].add(num)
                        match_order[name].append(numbers.index(num))

                total_matches = len(cumulative_matches[name])
                if total_matches == 5:
                    matched_names_5.add(name)
                    highlight_cell(name_cell, '#C0C0C0')
                elif total_matches == 6:
                    matched_names_6.add(name)

    if matched_names_6:
        earliest_last_match = None
        joint_winners = []
        other_six = []
        for nm in matched_names_6:
            if match_order[nm]:
                last_ix = max(match_order[nm])
                if earliest_last_match is None or last_ix < earliest_last_match:
                    earliest_last_match = last_ix
                    joint_winners = [nm]
                elif last_ix == earliest_last_match:
                    joint_winners.append(nm)
        for nm in matched_names_6:
            if nm not in joint_winners:
                other_six.append(nm)

        for table in doc.tables:
            if table.cell(0, 0).text in ["METADATA", "Date"]:
                continue
            for row in table.rows:
                if len(row.cells) < 2:
                    continue
                name_cell = row.cells[1]
                name = name_cell.text.strip()
                if name in joint_winners:
                    highlight_cell(name_cell, 'FFD700')
                elif name in other_six:
                    highlight_cell(name_cell, '#C0C0C0')
                elif name in matched_names_6:
                    clear_cell_highlight(name_cell)

        jackpot = calculate_jackpot(doc)
        winner_msg = generate_winner_message(joint_winners, other_six, jackpot)
        console_lines.append(f"Joint winners with 6: {joint_winners}")
        console_lines.append(winner_msg)
    else:
        console_lines.append("No entrants matched 6 yet.")

    if matched_names_5:
        console_lines.append(f"Names with exactly 5: {matched_names_5}")

    add_numbers_to_date_table(doc, numbers)

    return "\n".join(console_lines).strip()

###########################
# Kivy UI
###########################
# Note: The UI is defined in task.kv (see separate file)

from kivy.uix.scrollview import ScrollView

class NumberButton(Button):
    number = StringProperty()
    selected = BooleanProperty(False)

    def on_release(self):
        self.selected = not self.selected
        self.background_color = (1, 0.84, 0, 1) if self.selected else (0.3, 0.3, 0.3, 1)
        self.parent.parent.toggle_number(self.number, self.selected)

class MainScreen(BoxLayout):
    selected_numbers = ListProperty([])
    cumulative_matches = {}
    highlight_sequence = {}

    def __init__(self, **kwargs):
        super(MainScreen, self).__init__(**kwargs)
        download_file()
        self.build_number_grid()

    def build_number_grid(self):
        grid = self.ids.number_grid
        grid.clear_widgets()
        for i in range(1, 48):
            btn = NumberButton(
                text=str(i),
                number=str(i),
                background_color=(0.3, 0.3, 0.3, 1)
            )
            grid.add_widget(btn)

    def toggle_number(self, number, selected):
        if selected:
            self.selected_numbers.append(number)
        else:
            self.selected_numbers.remove(number)

    def highlight_selected(self):
        if not self.selected_numbers:
            self.show_popup("Error", "No numbers selected.")
            return

        doc = Document(LOCAL_FILE_PATH)
        console_msg = highlight_number(doc, self.selected_numbers, self.cumulative_matches, self.highlight_sequence)
        doc.save(LOCAL_FILE_PATH)
        upload_file()

        self.show_popup("Winner Logic", console_msg)

    def clear_selection(self):
        self.selected_numbers.clear()
        self.build_number_grid()
        self.show_popup("Cleared", "Selection cleared.")

    def reset_board(self):
        doc = Document(LOCAL_FILE_PATH)
        clear_highlights(doc, self.cumulative_matches, self.highlight_sequence)
        doc.save(LOCAL_FILE_PATH)
        upload_file()
        self.build_number_grid()
        self.show_popup("Reset", "All highlights + Date + METADATA removed.")

    def show_popup(self, title, text_content):
        scroll = ScrollView(size_hint=(1, 1))
        label = Label(
            text=text_content if text_content else "",
            color=(1,1,1,1),
            size_hint=(1, None),
            text_size=(300, None),
            halign="left",
            valign="top",
            padding=(10,10)
        )
        label.bind(
            texture_size=lambda instance, value: setattr(instance, 'height', value[1])
        )
        scroll.add_widget(label)

        popup = Popup(
            title=title,
            content=scroll,
            size_hint=(0.85, 0.7)
        )
        popup.open()

class BingoLottoApp(App):
    def build(self):
        self.title = "Bingo Lotto Highlighter"
        Builder.load_file("task.kv")
        return MainScreen()

if __name__ == '__main__':
    BingoLottoApp().run()
